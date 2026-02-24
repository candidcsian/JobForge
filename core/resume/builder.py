"""Simple resume builder from user input."""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_master_resume_from_manual(user_data, output_file):
    """Build master resume from manually entered data as Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Header - Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(user_data['name'].upper())
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_para = doc.add_paragraph()
    contact_text = f"{user_data['email']}"
    if user_data.get('phone'):
        contact_text += f" | {user_data['phone']}"
    if user_data.get('location'):
        contact_text += f" | {user_data['location']}"
    contact_para.add_run(contact_text)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Professional Experience
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
    exp_run.font.size = Pt(14)
    exp_run.font.bold = True
    
    # Add work history
    if 'manual_history' in user_data and user_data['manual_history']:
        for company in user_data['manual_history']:
            # Company and role
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{company['role']} | {company['company']}")
            company_run.font.bold = True
            company_run.font.size = Pt(11)
            
            # Duration
            duration_para = doc.add_paragraph()
            duration_para.add_run(f"{company['start_date']} - {company['end_date']}")
            duration_para.paragraph_format.space_after = Pt(6)
            
            # Technologies
            if company.get('technologies'):
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: {company['technologies']}")
                tech_run.font.italic = True
                tech_para.paragraph_format.space_after = Pt(6)
            
            # Responsibilities
            if company.get('responsibilities'):
                resp_heading = doc.add_paragraph()
                resp_heading.add_run('Key Responsibilities:').font.bold = True
                for resp in company['responsibilities']:
                    bullet = doc.add_paragraph(resp, style='List Bullet')
                    bullet.paragraph_format.left_indent = Pt(18)
            
            # Achievements
            if company.get('achievements'):
                ach_heading = doc.add_paragraph()
                ach_heading.add_run('Key Achievements:').font.bold = True
                for ach in company['achievements']:
                    bullet = doc.add_paragraph(ach, style='List Bullet')
                    bullet.paragraph_format.left_indent = Pt(18)
            
            doc.add_paragraph()  # Spacing between companies
    else:
        doc.add_paragraph("No work history provided.")
    
    # Save as Word document
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Change extension to .docx
    if output_path.suffix == '.md':
        output_path = output_path.with_suffix('.docx')
    
    doc.save(str(output_path))
    
    return output_path


def create_ats_resume_docx(user_data, output_file):
    """Create ATS-optimized resume in Word format."""
    
    doc = Document()
    
    # Header - Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(user_data.get('name', 'Your Name').upper())
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_para = doc.add_paragraph()
    contact_text = f"{user_data.get('email', 'your.email@example.com')}"
    if user_data.get('phone'):
        contact_text += f" | {user_data['phone']}"
    if user_data.get('location'):
        contact_text += f" | {user_data['location']}"
    contact_para.add_run(contact_text)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # If we have extracted text from uploaded resume, use that
    if user_data.get('extracted_text') and not user_data.get('manual_history'):
        # Add the extracted resume content
        content_para = doc.add_paragraph()
        content_para.add_run(user_data['extracted_text'])
        content_para.paragraph_format.space_after = Pt(12)
    
    # If we have manual history, use structured format
    elif user_data.get('manual_history'):
        # Professional Experience
        exp_heading = doc.add_paragraph()
        exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
        exp_run.font.size = Pt(14)
        exp_run.font.bold = True
        
        # Add work history
        for company in user_data['manual_history']:
            # Company and role
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{company['role']} | {company['company']}")
            company_run.font.bold = True
            company_run.font.size = Pt(11)
            
            # Duration
            duration_para = doc.add_paragraph()
            duration_para.add_run(f"{company['start_date']} - {company['end_date']}")
            duration_para.paragraph_format.space_after = Pt(6)
            
            # Responsibilities
            if company.get('responsibilities'):
                for resp in company['responsibilities']:
                    bullet = doc.add_paragraph(resp, style='List Bullet')
                    bullet.paragraph_format.left_indent = Pt(18)
            
            # Achievements
            if company.get('achievements'):
                for ach in company['achievements']:
                    bullet = doc.add_paragraph(ach, style='List Bullet')
                    bullet.paragraph_format.left_indent = Pt(18)
            
            doc.add_paragraph()  # Spacing
    else:
        # No content available
        doc.add_paragraph("No resume content available. Please provide your work history.")
    
    # Save
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    return output_path


def extract_skills_from_manual_history(user_data):
    """Extract skills from manually entered work history."""
    
    skills = set()
    
    if 'manual_history' in user_data:
        for company in user_data['manual_history']:
            # Extract from responsibilities
            for resp in company.get('responsibilities', []):
                # Simple keyword extraction
                words = resp.split()
                for word in words:
                    if len(word) > 3 and word[0].isupper():
                        skills.add(word.strip('.,;:'))
            
            # Extract from achievements
            for ach in company.get('achievements', []):
                words = ach.split()
                for word in words:
                    if len(word) > 3 and word[0].isupper():
                        skills.add(word.strip('.,;:'))
    
    return list(skills)


def create_skills_matrix(user_data, output_file):
    """Create skills matrix from user data."""
    
    skills = extract_skills_from_manual_history(user_data)
    
    content = f"""# {user_data['name']} - Skills Matrix

Generated: {datetime.now().strftime('%Y-%m-%d')}

---

## Technical Skills

"""
    
    for skill in sorted(skills):
        content += f"- {skill}\n"
    
    content += f"\n\n**Total Skills Identified:** {len(skills)}\n"
    
    # Write to file
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content)
    
    return output_path
